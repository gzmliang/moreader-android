#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PROJECT_DIR = "/root/.hermes/cache/projects/moreader-android"

def create_source_code_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    title = doc.add_heading('墨阅 Moreader — 源代码文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for label, value in [
        ("软件名称", "墨阅 Moreader EPUB 阅读器"),
        ("软件版本", "v1.3.3"),
        ("开发语言", "Kotlin"),
        ("开发框架", "Jetpack Compose + Android SDK"),
        ("代码行数", "约 5,800 行"),
        ("文件数量", "32 个源文件"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}：")
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
    
    doc.add_page_break()
    
    kt_files = []
    base = os.path.join(PROJECT_DIR, "app/src/main/java")
    for root, dirs, files in os.walk(base):
        for f in files:
            if f.endswith(".kt"):
                full = os.path.join(root, f)
                rel = os.path.relpath(full, base)
                with open(full, "r", encoding="utf-8") as fh:
                    content = fh.read()
                kt_files.append((rel, content))
    kt_files.sort()
    
    for rel_path, content in kt_files:
        doc.add_heading(rel_path, level=1)
        code_para = doc.add_paragraph()
        run = code_para.add_run(content)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        doc.add_page_break()
    
    output = os.path.join(PROJECT_DIR, "墨阅Moreader_源代码.docx")
    doc.save(output)
    print(f"OK: {output}")

def create_manual_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Title page
    title = doc.add_heading('墨阅 Moreader 用户使用说明书', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for label, value in [
        ("软件名称", "墨阅 Moreader EPUB 阅读器"),
        ("软件版本", "v1.3.3"),
        ("适用平台", "Android 8.0 及以上"),
        ("开发者", "梁竞敏"),
        ("编写日期", "2026年5月"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}：")
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
    
    doc.add_page_break()
    
    # Table of contents
    doc.add_heading('目录', level=1)
    toc_items = [
        "1. 软件简介",
        "2. 安装与启动",
        "3. 主界面说明",
        "4. 阅读功能",
        "   4.1 打开书籍",
        "   4.2 翻页操作",
        "   4.3 目录导航",
        "   4.4 阅读进度",
        "5. 主题与显示设置",
        "   5.1 阅读主题",
        "   5.2 字体大小",
        "6. 文字选择与工具栏",
        "   6.1 选中文本",
        "   6.2 加入生词本",
        "   6.3 高亮标注",
        "   6.4 朗读选中文字",
        "   6.5 AI翻译",
        "   6.6 AI解释",
        "   6.7 语法分析",
        "7. 生词本功能",
        "   7.1 查看生词",
        "   7.2 生词释义",
        "   7.3 生词发音",
        "   7.4 重新拉取释义",
        "   7.5 生词导出",
        "   7.6 删除生词",
        "8. TTS 朗读设置",
        "   8.1 TTS 引擎选择",
        "   8.2 语速调节",
        "   8.3 朗读控制",
        "9. 书签功能",
        "10. 高级设置",
        "    10.1 AI 服务配置",
        "    10.2 TTS 服务配置",
        "11. 常见问题",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # Chapter 1
    doc.add_heading('1. 软件简介', level=1)
    doc.add_paragraph(
        '墨阅 Moreader 是一款专为 Android 设备设计的专业 EPUB 电子书阅读器。'
        '软件采用 Kotlin + Jetpack Compose 技术栈开发，提供流畅的阅读体验和强大的学习辅助功能。'
        '主要特色包括：'
    )
    features = [
        '精美的 EPUB 电子书渲染，支持多种排版格式',
        'AI 驱动的生词本，支持中英文双语结构化释义',
        '多引擎 TTS 朗读（Edge TTS / AI Voice / 自定义 TTS）',
        '丰富的阅读主题（白天、夜间、护眼等模式）',
        '文本高亮标注与书签管理',
        '选区智能翻译、解释与语法分析',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    # Chapter 2
    doc.add_heading('2. 安装与启动', level=1)
    doc.add_heading('2.1 安装', level=2)
    doc.add_paragraph(
        '下载墨阅 Moreader APK 安装包后，在 Android 设备上进行安装。'
        '如果系统提示"禁止安装未知来源应用"，请在设置中允许来自该来源的安装。'
    )
    doc.add_heading('2.2 启动', level=2)
    doc.add_paragraph(
        '安装完成后，点击桌面"墨阅 Moreader"图标即可启动应用。首次启动时，'
        '应用会显示书架界面，等待用户导入 EPUB 电子书文件。'
    )
    
    # Chapter 3
    doc.add_heading('3. 主界面说明', level=1)
    doc.add_paragraph(
        '启动应用后进入书架界面。书架以卡片形式展示已导入的书籍，'
        '每张卡片显示书籍封面、书名和阅读进度。点击任意书籍卡片即可开始阅读。'
        '书架界面右上角设有设置按钮，可进入全局设置页面。'
    )
    
    # Chapter 4
    doc.add_heading('4. 阅读功能', level=1)
    doc.add_heading('4.1 打开书籍', level=2)
    doc.add_paragraph(
        '在书架界面点击任意书籍即可打开。应用会自动加载书籍内容并恢复上次的阅读进度。'
    )
    doc.add_heading('4.2 翻页操作', level=2)
    doc.add_paragraph(
        '阅读界面支持左右滑动翻页。向左滑动翻到下一页，向右滑动翻到上一页。'
        '点击屏幕中央区域可切换顶部工具栏和底部导航栏的显示/隐藏。'
    )
    doc.add_heading('4.3 目录导航', level=2)
    doc.add_paragraph(
        '点击顶部工具栏的目录图标，可打开章节目录面板。'
        '在目录面板中选择任意章节即可跳转到该章节阅读。'
    )
    doc.add_heading('4.4 阅读进度', level=2)
    doc.add_paragraph(
        '底部导航栏显示阅读进度条，可拖动进度条跳转到任意章节。'
        '进度条左侧显示当前章节序号和总章节数，右侧显示阅读进度百分比。'
    )
    
    # Chapter 5
    doc.add_heading('5. 主题与显示设置', level=1)
    doc.add_heading('5.1 阅读主题', level=2)
    doc.add_paragraph(
        '在阅读界面点击底部工具栏的主题按钮，可切换不同的阅读主题。'
        '支持白天模式、夜间模式、护眼模式等多种配色方案，'
        '满足不同环境下的阅读需求。'
    )
    doc.add_heading('5.2 字体大小', level=2)
    doc.add_paragraph(
        '在底部导航栏中，点击字体大小的加减按钮可调节文字大小。'
        '字体大小范围为 14 到 28 磅，可根据个人阅读习惯自由调整。'
    )
    
    # Chapter 6
    doc.add_heading('6. 文字选择与工具栏', level=1)
    doc.add_heading('6.1 选中文本', level=2)
    doc.add_paragraph(
        '长按阅读界面中的任意文字，系统会自动选中该文字所在的段落或词汇。'
        '拖动选区手柄可调整选中范围。松开手指后，底部会出现工具栏。'
    )
    doc.add_heading('6.2 加入生词本', level=2)
    doc.add_paragraph(
        '点击工具栏的"📖加入生词本"按钮，即可将选中的文字添加到生词本。'
        '系统会自动调用 AI 服务查询该词语的释义。如果该词已在生词本中，会提示"该词已存在"。'
    )
    doc.add_heading('6.3 高亮标注', level=2)
    doc.add_paragraph(
        '点击工具栏的高亮按钮，可为选中的文字添加黄色高亮标记。'
        '已高亮的文字再次选中时，会显示"取消高亮"选项。'
    )
    doc.add_heading('6.4 朗读选中文字', level=2)
    doc.add_paragraph(
        '点击工具栏的"🔊朗读"按钮，系统会使用当前配置的 TTS 引擎朗读选中的文字。'
    )
    doc.add_heading('6.5 AI翻译', level=2)
    doc.add_paragraph(
        '点击工具栏的"翻译"按钮，系统会调用 AI 服务对选中文字进行翻译。'
        '翻译结果会在弹窗中显示。自动识别中文或英文输入，输出对应的双语翻译。'
    )
    doc.add_heading('6.6 AI解释', level=2)
    doc.add_paragraph(
        '点击工具栏的"解释"按钮，系统会提供选中文字的语境解析、'
        '重点词汇分析和用法提示，帮助深入理解文本内容。'
    )
    doc.add_heading('6.7 语法分析', level=2)
    doc.add_paragraph(
        '点击工具栏的"语法分析"按钮，系统会对选中的句子进行结构分析，'
        '包括主谓宾定状补、语法要点和易错提示。'
    )
    
    # Chapter 7
    doc.add_heading('7. 生词本功能', level=1)
    doc.add_heading('7.1 查看生词', level=2)
    doc.add_paragraph(
        '在主界面点击生词本图标，进入生词列表页面。列表按添加时间倒序排列，'
        '每个词条显示原词/原字及简要释义。'
    )
    doc.add_heading('7.2 生词释义', level=2)
    doc.add_paragraph(
        '点击生词列表中的任意词条，可展开查看详细释义。'
        '释义包含以下内容：'
    )
    detail_items = [
        '原词/原字（大字显示）',
        '拼音或音标',
        '词性标注',
        '中文释义（1-2条）',
        '英文释义（English Definition，1-2条）',
        '组词/派生词（2个常见搭配）',
        '双语例句（中英文对照，1条）',
    ]
    for item in detail_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.3 生词发音', level=2)
    doc.add_paragraph(
        '在生词详情页面，点击词条右侧的🔊喇叭图标，即可听取该词的发音。'
        '发音引擎与阅读设置中配置的 TTS 服务一致。'
    )
    doc.add_heading('7.4 重新拉取释义', level=2)
    doc.add_paragraph(
        '如果某条生词的释义不完整或显示异常，点击词条下方的'
        '"🔄 重新拉取释义"按钮，系统会重新调用 AI 服务获取最新的结构化释义。'
    )
    doc.add_heading('7.5 生词导出', level=2)
    doc.add_paragraph(
        '在生词本页面，支持将生词列表导出为 Markdown 或 CSV 格式文件，'
        '方便在其他应用或设备中使用。'
    )
    doc.add_heading('7.6 删除生词', level=2)
    doc.add_paragraph(
        '在生词详情页面，点击🗑️垃圾桶图标即可删除该词条。删除操作不可恢复，请谨慎操作。'
    )
    
    # Chapter 8
    doc.add_heading('8. TTS 朗读设置', level=1)
    doc.add_heading('8.1 TTS 引擎选择', level=2)
    doc.add_paragraph(
        '墨阅 Moreader 支持多种 TTS 朗读引擎，可在设置页面进行切换：'
    )
    tts_items = [
        'Edge TTS：微软 Edge 浏览器的在线 TTS 服务，音质优秀，支持多种语音',
        'AI Voice：基于 AI 模型的 TTS 服务，支持自定义模型和语音',
        '自定义 TTS：兼容 OpenAI 格式的 TTS 服务接口',
    ]
    for item in tts_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('8.2 语速调节', level=2)
    doc.add_paragraph(
        '在 TTS 设置面板中，可通过滑块调节朗读语速。'
        '语速范围为 0.5 到 2.0 倍速，默认 1.0 倍速。'
    )
    doc.add_heading('8.3 朗读控制', level=2)
    doc.add_paragraph(
        '底部导航栏提供朗读控制按钮：播放/暂停、停止。'
        '点击播放后，系统会从当前章节开始逐段朗读全书内容，'
        '并在阅读界面高亮当前朗读的段落。'
    )
    
    # Chapter 9
    doc.add_heading('9. 书签功能', level=1)
    doc.add_paragraph(
        '在阅读界面点击顶部工具栏的书签图标，可为当前阅读位置添加书签。'
        '书签列表可在主界面或阅读界面的书签面板中查看和管理。'
        '点击书签即可跳转到对应位置，支持删除不需要的书签。'
    )
    
    # Chapter 10
    doc.add_heading('10. 高级设置', level=1)
    doc.add_heading('10.1 AI 服务配置', level=2)
    doc.add_paragraph(
        '生词本查词、AI翻译、AI解释、语法分析等功能均需要配置 AI 服务。'
        '在设置页面填写以下信息：'
    )
    ai_items = [
        '服务端点（Endpoint）：OpenAI 兼容的 API 地址，如 https://api.example.com/v1',
        'API Key：服务访问密钥',
        '模型名称（Model）：使用的 AI 模型，如 gpt-3.5-turbo',
    ]
    for item in ai_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.2 TTS 服务配置', level=2)
    doc.add_paragraph(
        '根据选择的 TTS 引擎类型，需要填写对应的服务配置信息：'
    )
    tts_config_items = [
        'Edge TTS：服务端点、语音名称',
        'AI Voice：服务端点、API Key、模型名称、语音 ID',
        '自定义 TTS：服务端点、API Key、模型名称、语音名称',
    ]
    for item in tts_config_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Chapter 11
    doc.add_heading('11. 常见问题', level=1)
    
    faqs = [
        ("Q: 为什么生词本查词失败？",
         "A: 请检查设置页面是否正确配置了 AI 服务的 Endpoint 和 API Key。确认无误后重试。"),
        ("Q: 为什么朗读没有声音？",
         "A: 请检查 TTS 设置中的服务配置是否正确，网络连接是否正常。可切换不同 TTS 引擎尝试。"),
        ("Q: 如何导入 EPUB 书籍？",
         "A: 在书架界面点击加号按钮，从设备存储中选择 EPUB 文件即可导入。"),
        ("Q: 书籍的字体和排版可以调整吗？",
         "A: 可以在阅读界面中调整字体大小和阅读主题。EPUB 书籍原有的排版样式会尽量保留。"),
        ("Q: 生词数据会丢失吗？",
         "A: 生词本数据保存在本地数据库中。卸载应用会导致数据丢失，建议定期导出生词列表作为备份。"),
        ("Q: 支持哪些 Android 版本？",
         "A: 墨阅 Moreader 支持 Android 8.0（API 26）及以上的设备。"),
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        r1 = p.add_run(q)
        r1.bold = True
        p2 = doc.add_paragraph(a)
        p2.paragraph_format.space_after = Pt(12)
    
    output = os.path.join(PROJECT_DIR, "墨阅Moreader_使用说明书.docx")
    doc.save(output)
    print(f"OK: {output}")

create_source_code_doc()
create_manual_doc()
